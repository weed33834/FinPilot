"""Word 文档解析器：基于 python-docx 提取段落文本与表格。

支持 .docx；旧版 .doc（OLE 二进制）python-docx 无法解析，会抛出 ParserError。
"""
from __future__ import annotations

import os
import zipfile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from .base import BaseParser, ParsedDocument, ParsedPage, ParserError


class DocxParser(BaseParser):
    """解析 Word 文档，提取段落文本与表格，整篇作为一页。"""

    def parse(self, file_path: str) -> ParsedDocument:
        filename = self._check_file(file_path)
        ext = os.path.splitext(file_path)[1].lower().lstrip(".")

        # python-docx 仅支持 .docx（OOXML zip 包），无法解析旧版 .doc
        if ext == "doc":
            raise ParserError(
                f"不支持旧版 .doc 格式 [{filename}]，请转换为 .docx 后再解析"
            )

        try:
            doc = Document(file_path)
        except (OSError, PackageNotFoundError, zipfile.BadZipFile, ValueError) as e:
            # PackageNotFoundError/BadZipFile: 非 docx 或文件损坏
            raise ParserError(f"Word 解析失败 [{filename}]: {e}") from e

        # 收集正文段落文本
        text_parts: list[str] = [para.text for para in doc.paragraphs]
        # 提取全部表格，转为二维字符串列表
        tables: list[list[list[str]]] = [
            self._table_to_list(tbl) for tbl in doc.tables
        ]

        page = ParsedPage(
            page_number=1,
            text="\n".join(text_parts),
            tables=tables,
        )

        return ParsedDocument(
            filename=filename,
            file_type="docx",
            pages=[page],
            tables=tables,
            metadata={
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            },
        )

    @staticmethod
    def _table_to_list(table) -> list[list[str]]:
        """python-docx 表格转为二维字符串列表。"""
        return [[cell.text for cell in row.cells] for row in table.rows]
